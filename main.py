import requests
from docx import Document
import os

input_path = "input/letra_musica.docx"
output_path = "output/letra_musica_traduzida.docx"

api_key = "SUA_CHAVE_DE_API"
endpoint = "SEU_ENDPOINT"
translator_endpoint = f"{endpoint}/translate?api-version=3.0"

def read_word_file(file_path):
   doc = Document(file_path)
   return "\n".join([para.text for para in doc.paragraphs])

def write_word_file(text, file_path):
   doc = Document()
   doc.add_paragraph(text)
   doc.save(file_path)

def translate_text(text, to_language="pt-br"):
   headers = {
       "Ocp-Apim-Subscription-Key": api_key,
       "Content-Type": "application/json",
       "Ocp-Apim-Subscription-Region": "REGIÃO_DO_SEU_SERVIÇO"
   }
   body = [{"text": text}]
   params = {"to": to_language}

   response = requests.post(translator_endpoint, headers=headers, json=body, params=params)
   response.raise_for_status()

   translation = response.json()[0]["translations"][0]["text"]
   return translation

try:
   original_text = read_word_file(input_path)

   translated_text = translate_text(original_text)

   write_word_file(translated_text, output_path)

   print("Tradução realizada com sucesso!")
except Exception as e:
   print(f"Erro ao realizar a tradução: {e}")